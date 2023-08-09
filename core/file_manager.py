"""
A module for file manager in the core package.
"""
import logging
import os
from abc import ABC, abstractmethod
from enum import Enum
from pathlib import Path
from typing import Any, Optional, Union

import pandas as pd
from docx import Document

logger: logging.Logger = logging.getLogger(__name__)


class DataType(str, Enum):
    """
    Data Type class based on Enum
    """

    RAW: str = "data/raw/"
    PROCESSED: str = "data/processed/"
    FIGURES: str = "reports/figures/"


class FileManager(ABC):
    """
    Abstract class for file management.
    """

    @abstractmethod
    def load(
        self, filename: Union[str, Path], data_type: Optional[DataType]
    ) -> pd.DataFrame:
        """
        Abstract method to load data from file.
        :param filename: The name of the file, including extension.
        :type filename: Union[str, Path]
        :param data_type: The path where data will be saved.
        :type data_type: Optional[DataType]
        :return: Dataframe retrieved from file.
        :rtype: pd.DataFrame
        """

    @abstractmethod
    def save(
        self,
        dataframe: pd.DataFrame,
        data_type: Optional[DataType],
        filename: str,
    ) -> bool:
        """
        Abstract method to save data to file.
        :param dataframe: DataFrame to save.
        :type dataframe: pd.DataFrame
        :param data_type: Path where data will be saved.
        :type data_type: Optional[DataType]
        :param filename: Name of the file.
        :type filename: str
        :return: True if the file was created; otherwise false.
        :rtype: bool
        """


class XLSXManager(FileManager):
    def load(
        self, filename: Union[str, Path], data_type: Optional[DataType]
    ) -> pd.DataFrame:
        if data_type:
            filename = os.path.join(data_type, filename)
        return pd.read_excel(filename, engine="openpyxl")

    def save(
        self,
        dataframe: pd.DataFrame,
        data_type: Optional[DataType],
        filename: str,
    ) -> bool:
        if len(dataframe) == 0:
            return False
        if data_type:
            filename = os.path.join(data_type, filename)
        with pd.ExcelWriter(filename, engine="openpyxl") as writer:
            dataframe.to_excel(writer, index=False)
        return True


class CSVManager(FileManager):
    def load(
        self, filename: Union[str, Path], data_type: Optional[DataType]
    ) -> pd.DataFrame:
        if data_type:
            filename = os.path.join(data_type, filename)
        return pd.read_csv(filename)

    def save(
        self,
        dataframe: pd.DataFrame,
        data_type: Optional[DataType],
        filename: str,
    ) -> bool:
        if len(dataframe) == 0:
            return False
        if data_type:
            filename = os.path.join(data_type, filename)
        dataframe.to_csv(filename, index=False)
        return True


class DOCXManager(FileManager):
    def load(
        self, filename: Union[str, Path], data_type: Optional[DataType]
    ) -> pd.DataFrame:
        if data_type:
            filename = os.path.join(data_type, filename)
        doc = Document(filename)
        text: list[Any] = [p.text for p in doc.paragraphs]
        return pd.DataFrame(text)

    def save(
        self,
        dataframe: pd.DataFrame,
        data_type: Optional[DataType],
        filename: str,
    ) -> bool:
        if len(dataframe) == 0:
            return False
        if data_type:
            filename = os.path.join(data_type, filename)
        doc = Document()
        for index, row in dataframe.iterrows():
            doc.add_paragraph(str(row))
        doc.save(filename)
        return True
